import sys
import os
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

import anthropic
import pandas as pd
from docx import Document
from config import ANTHROPIC_API_KEY


def extract_cv_text(cv_path: str) -> str:
    """Extract plain text from a .docx CV file."""
    doc = Document(cv_path)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    return "\n".join(paragraphs)


def tailor_with_claude(client: anthropic.Anthropic, cv_text: str, job: dict) -> dict:
    """
    Send CV + job description to Claude.
    Returns dict with cover_letter and ats_answers.
    """
    title   = job.get("title", "")
    company = job.get("company", "")
    desc    = job.get("description", "")[:3000]

    prompt = f"""You are helping a real person write job application materials. Your output must read like a human wrote it. Follow every rule below strictly.

WRITING RULES:

ABSOLUTE RULE: Never use em dashes anywhere in your output. Use commas, semicolons, or periods instead. Not even one em dash. This is non-negotiable.

Banned words — never use any of these:
delve, dive into, navigate (figurative), underscore, bolster, foster, harness, leverage, unpack, shed light on, pave the way, pivotal, groundbreaking, cutting-edge, transformative, game-changing, innovative, robust, comprehensive, seamless, intricate, nuanced (as empty praise), vibrant, multifaceted, holistic, testament, landscape (figurative), realm, spearheaded, orchestrated, synergised, utilised, passionate, dynamic

Banned phrases — never use:
- "In today's [fast-paced/rapidly evolving/digital] world..."
- "It's important/worth noting that..."
- "One of the most [important/significant/crucial]..."
- "When it comes to..." / "At its core..." / "At the end of the day..."
- "This is where X comes in" / "Let's break it down"
- "Plays a crucial role in..." / "It cannot be overstated..."
- "I am writing to express my interest..."
- "I am excited to apply..."

Banned structures:
- "It's not just X, it's Y"
- "Not only X, but Y"
- "This isn't about X. It's about Y."
- "No X. No Y. Just Z."

Style rules:
- Use contractions. "It's," "don't," "won't," "I've."
- NO EM DASHES anywhere. Use commas or parentheses instead.
- Vary sentence and paragraph length. Mix short punchy lines with longer ones.
- No signposting ("Let's explore," "Now let's turn to"). Just make your point.
- Don't open with a sweeping contextual statement. Start on substance.
- Don't close with a summary or inspirational wrap-up. End on substance.
- No performative enthusiasm ("exciting," "incredible," "powerful").
- No preamble. Get straight to the content.
- Never repeat the same point in different words. Say it once.
- Read it out loud test: if any sentence sounds like a press release, rewrite it.

CANDIDATE'S CV:
{cv_text}

JOB THEY ARE APPLYING FOR:
Title: {title}
Company: {company}
Description:
{desc}

Produce exactly TWO sections separated by the markers shown. No text before the first marker.

=== COVER LETTER ===
Write a cover letter for this role. Maximum 250 words.
Rules:
- Start with something specific about the company or role, not a generic opener
- Reference 2-3 specific skills from the CV that directly match this job
- Include one brief concrete example from their experience that shows relevant impact
- End with a clear call to action, confident not desperate
- Use contractions throughout, write like you are talking to a hiring manager
- Don't use "Dear Hiring Manager" — use "Hi [Company] team," or address by company name
- Don't restate the job title in your opening sentence, they know what they posted
- NO EM DASHES anywhere in the letter

=== ATS ANSWERS ===
Pre-filled answers for common application form questions. Be specific and human, not template-sounding.

1. Tell us about yourself:
(2-3 sentences tailored to this role, conversational not rehearsed)

2. Why do you want to work at {company}?:
(2-3 sentences, specific to what this company actually does, not generic flattery)

3. What relevant experience do you have?:
(3-4 sentences highlighting the best matches between the CV and this specific job)

4. Expected salary range:
Open to discussion based on the full compensation package

5. Work authorisation:
Please discuss during interview

6. Notice period:
Available to discuss

7. Key skills for this role:
(comma-separated list of 8-10 skills pulled directly from the CV that match this job, no invented skills)
"""

    try:
        response = client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=2000,
            messages=[{"role": "user", "content": prompt}],
        )
        full_text = response.content[0].text

        # Post-process: strip any em dashes that slipped through
        full_text = full_text.replace("—", ",").replace("\u2014", ",")

        sections = {
            "cover_letter": "",
            "ats_answers":  "",
        }

        markers = [
            ("=== COVER LETTER ===", "cover_letter"),
            ("=== ATS ANSWERS ===",  "ats_answers"),
        ]

        for i, (marker, key) in enumerate(markers):
            if marker not in full_text:
                continue
            start = full_text.index(marker) + len(marker)
            end   = len(full_text)
            for next_marker, _ in markers[i+1:]:
                if next_marker in full_text:
                    next_idx = full_text.index(next_marker)
                    if next_idx > start:
                        end = next_idx
                        break
            sections[key] = full_text[start:end].strip()

        return sections

    except Exception as e:
        print(f"    -> Claude API error: {e}")
        return {"cover_letter": "", "ats_answers": "", "error": str(e)}


def save_docx(content: str, filepath: str):
    """Save plain text content to a .docx file with clean formatting."""
    doc = Document()
    for line in content.split("\n"):
        doc.add_paragraph(line)
    doc.save(filepath)


def run_tailor(scored_results: dict) -> dict:
    """
    Generate cover letter and ATS answers for every matched job.
    Input:  dict of {campaign_name: scored_df}
    Output: dict of {campaign_name: df_with_file_paths}
    """
    from config import CAMPAIGNS
    print("\n=== MODULE 4: AI tailoring with Claude ===")

    if not ANTHROPIC_API_KEY:
        print("  ERROR: ANTHROPIC_API_KEY not found in .env — check your .env file")
        return {}

    client   = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)
    results  = {}
    cv_cache = {}

    for campaign_name, df in scored_results.items():
        campaign = CAMPAIGNS.get(campaign_name, {})
        cv_path  = campaign.get("cv", "")

        print(f"\n  [{campaign_name}] — {len(df)} jobs to tailor")

        if df.empty:
            continue

        # Load and cache CV text
        if cv_path not in cv_cache:
            try:
                cv_cache[cv_path] = extract_cv_text(cv_path)
                print(f"  CV loaded: {cv_path}")
            except Exception as e:
                print(f"  ERROR loading CV {cv_path}: {e}")
                continue

        cv_text = cv_cache[cv_path]

        df = df.copy()
        df["cover_letter_path"] = ""
        df["ats_answers_path"]  = ""

        for idx, row in df.iterrows():
            # Build a clean folder name from company + title
            company_safe = "".join(
                c for c in row["company"] if c.isalnum() or c in " _-"
            )[:30].strip().replace(" ", "_")
            title_safe = "".join(
                c for c in row["title"] if c.isalnum() or c in " _-"
            )[:30].strip().replace(" ", "_")
            folder = os.path.join("output", f"{company_safe}_{title_safe}")
            os.makedirs(folder, exist_ok=True)

            print(f"\n    [{idx+1}/{len(df)}] {row['title']} @ {row['company']}")
            print(f"    Score: {row['match_score']:.3f} | Source: {row['source']}")

            result = tailor_with_claude(client, cv_text, row.to_dict())

            if result.get("error"):
                print(f"    -> Failed: {result['error']}")
                continue

            # Save cover letter
            if result["cover_letter"]:
                cl_file = os.path.join(folder, "cover_letter.docx")
                save_docx(result["cover_letter"], cl_file)
                df.at[idx, "cover_letter_path"] = cl_file
                print(f"    -> Cover letter: {cl_file}")
            else:
                print(f"    -> WARNING: Cover letter was empty")

            # Save ATS answers
            if result["ats_answers"]:
                ats_file = os.path.join(folder, "ats_answers.docx")
                save_docx(result["ats_answers"], ats_file)
                df.at[idx, "ats_answers_path"] = ats_file
                print(f"    -> ATS answers:  {ats_file}")
            else:
                print(f"    -> WARNING: ATS answers were empty")

        results[campaign_name] = df

    total = sum(len(df) for df in results.values())
    print(f"\n=== Module 4 complete — {total} jobs tailored across "
          f"{len(results)} campaigns ===")
    return results


if __name__ == "__main__":
    from modules.scraper import run_scraper
    from modules.cleaner import run_cleaner
    from modules.matcher import run_matcher

    scraped  = run_scraper()
    cleaned  = run_cleaner(scraped)
    scored   = run_matcher(cleaned)
    tailored = run_tailor(scored)

    print("\n--- Generated files ---")
    for campaign_name, df in tailored.items():
        print(f"\n[{campaign_name}]")
        for _, row in df.iterrows():
            print(f"  {row['title']} @ {row['company']}")
            print(f"    Cover letter: {row['cover_letter_path']}")
            print(f"    ATS answers:  {row['ats_answers_path']}")

    print("\n✓ Module 4 test complete")